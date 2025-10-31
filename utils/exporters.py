from typing import Dict
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def export_to_word(summary, output_path):
    """
    Exports the structured summary to a Word document.

    :param summary: The summary dictionary.
    :param output_path: Path to save the Word file.
    """
    doc = Document()
    doc.add_heading('Research Paper Summary', 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(summary)
    # for section, content in summary.items():
    #     heading = doc.add_heading(section.capitalize(), level=1)
    #     heading.runs[0].font.size = Pt(14)
    #     p = doc.add_paragraph(content)
    #     p.paragraph_format.space_after = Pt(12)

    doc.save(output_path)
    print(f"Summary exported to: {output_path}")